import os
import requests
import sys
import subprocess
from docx import Document
import openai
import webbrowser
from pytube import YouTube
import uuid
from tqdm import tqdm  # Import tqdm for progress bars
from rich.progress import Progress  # Import Progress for a more stylish output
import json
from slugify import slugify


if len(sys.argv) >= 4:
    custom_output_filename = sys.argv[3]
else:
    custom_output_filename = random_uuid = str(uuid.uuid4())

# Generate filenames with different extensions
video_filename = os.path.join('downloads', f'{custom_output_filename}.mp4')
audio_filename = os.path.join('downloads', f'{custom_output_filename}.mp3')
#document_filename = os.path.join('downloads', f'{custom_output_filename}.docx')
#json_filename= os.path.join('downloads', f'{custom_output_filename}.json')

class AudioTranscriptionService:
    def __init__(self, model="whisper-1"):
        self.model = model

    def transcribe(self, audio_file_path):
        with tqdm(total=100, unit="%", desc="Transcribing Audio") as progress_bar:
            try:
                with open(audio_file_path, 'rb') as audio_file:
                    audio_content = audio_file.read()
                    transcription = openai.audio.transcriptions.create(
                        model=self.model,
                        file=audio_file
                    )                
                    return transcription.text

            finally:
                progress_bar.update(100)


class MeetingAnalyzer:
    def __init__(self, transcription):
        self.transcription = transcription

    def summarize(self):
        return self._get_openai_response("summarize into a concise abstract paragraph")

    def extract_key_points(self):
        return self._get_openai_response("identify and list the main points")

    def extract_action_items(self):
        return self._get_openai_response("extract action items")

    def analyze_sentiment(self):
        return self._get_openai_response("analyze the sentiment of the following text", message_key='message')
    
    def gen_title(self):
        return self._get_openai_response("generate a title for the transcription.")

    def _get_openai_response(self, task_description, message_key='message.content'):
        response = openai.chat.completions.create(
            model="gpt-4-1106-preview",
            temperature=0,
            messages=[
                {"role": "system", "content": f"You are an AI {task_description}."},
                {"role": "user", "content": self.transcription}
            ]
        )
        return response.choices[0].message.content


class DocumentManager:
    @staticmethod
    def save_to_docx(content, filename):
        doc = Document()
        for key, value in content.items():
            heading = key.replace('_', ' ').title()
            doc.add_heading(heading, level=1)
            doc.add_paragraph(value)
            doc.add_paragraph()
        doc.save(filename)


class VideoDownloader:
    @staticmethod
    def download_from_google_drive(drive_url):
        file_id = drive_url.split('/')[-2]
        direct_download_url = f'https://drive.google.com/uc?id={file_id}&export=download'
        with requests.get(direct_download_url, stream=True) as r:
            r.raise_for_status()
            with open(video_filename, 'wb') as f:
                for chunk in r.iter_content(chunk_size=8192):
                    f.write(chunk)
        return video_filename

    @staticmethod
    def download_from_youtube(youtube_url):
        yt = YouTube(youtube_url)
        yt = yt.streams.get_highest_resolution()
        try:
            yt.download(filename=video_filename)
        except:
            print("Hubo un error al descargar el video del URL proporcionado...")

        return video_filename

    @staticmethod
    def autodetect_download_source(url):
        with tqdm(total=100, unit="%", desc="Downloading video") as progress_bar:
            try:
                if 'drive.google.com' in url:
                    return VideoDownloader.download_from_google_drive(url)
                elif 'youtube.com' in url:
                    return VideoDownloader.download_from_youtube(url)
                else:
                    return url  # assume local file
            except subprocess.CalledProcessError as e:
                print(f"Error during audio extraction: {e.stderr.decode()}")
                return None
            finally:
                progress_bar.update(100)


class AudioExtractor:
    @staticmethod
    def extract_audio(input_video, bitrate='56k'):
        if os.path.exists(audio_filename):
            answer = input(
                "Output audio file already exists. Delete it? (yes/no): ")
            if answer.lower() == 'yes':
                print(f"Deleting {audio_filename}...")
                os.remove(audio_filename)
            else:
                return audio_filename

        # Use tqdm to display a progress bar during audio extraction
        with tqdm(total=100, unit="%", desc="Extracting audio") as progress_bar:
            try:
                subprocess.run([
                    'ffmpeg',
                    '-i', input_video,
                    '-b:a', bitrate,
                    '-vn',
                    audio_filename
                ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.PIPE)
            except subprocess.CalledProcessError as e:
                print(f"Error during audio extraction: {e.stderr.decode()}")
                return None
            finally:
                progress_bar.update(100)
        
        return audio_filename


def login_with_google():
    webbrowser.open('https://accounts.google.com/o/oauth2/auth?client_id=YOUR_CLIENT_ID&redirect_uri=YOUR_REDIRECT_URI&scope=https://www.googleapis.com/auth/drive.readonly&response_type=code', new=2)


def main(api_key, file_path):

    print("Beginning video transcription process for run %s" %
          (custom_output_filename))

    video_file = VideoDownloader.autodetect_download_source(
            file_path)
    # Use tqdm to display a progress bar while extracting audio
    audio_file = AudioExtractor.extract_audio(video_file)
    transcription_text = AudioTranscriptionService().transcribe(audio_file)
    analyzer = MeetingAnalyzer(transcription_text)
    
    meeting_info = {
        'meeting_title': analyzer.gen_title(),
        'abstract_summary': analyzer.summarize(),
        'key_points': analyzer.extract_key_points(),
        'action_items': analyzer.extract_action_items(),
        'sentiment': analyzer.analyze_sentiment()
    }

    document_filename = os.path.join('downloads', f'{slugify(meeting_info["meeting_title"])}.docx')
    json_filename= os.path.join('downloads', f'{slugify(meeting_info["meeting_title"])}.json')

    # Save the meeting information to a JSON file
    with open(json_filename, 'w') as json_file:
        json.dump(meeting_info, json_file, indent=4)

    DocumentManager.save_to_docx(meeting_info, document_filename)
    print("Document saved: %s" % document_filename)

    print("Do you want to see the meeting information? (yes/no): ")
    show_info = input()
    if show_info.lower() == 'yes':
        print("\nMeeting Information:")
        for section, content in meeting_info.items():
            print(f"{section.title().replace('_', ' ')}:\n{content}\n")

    # After analyzing the meeting, add the following code to delete the intermediate video and audio files:
    if os.path.exists(video_filename):
        os.remove(video_filename)
        print(f"Deleted intermediate video file: {video_filename}")

    if os.path.exists(audio_filename):
        os.remove(audio_filename)
        print(f"Deleted intermediate audio file: {audio_filename}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python transcribe_video.py api_key file_path")
        sys.exit(1)

    main(sys.argv[1], sys.argv[2])
